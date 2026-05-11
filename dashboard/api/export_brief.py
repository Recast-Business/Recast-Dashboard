"""POST /api/export_brief — Build a creator-assessment brief as a Google Doc.

Body:
  {
    "partner": "Brand or agency name",
    "campaign": "Campaign title",
    "criteria": "Free-form description of requirements, geos, CCV range, etc.",
    "month_year": "April 2026",          # optional, defaults to today
    "creator_ids": ["uuid", "uuid", ...], # from Leads selection
  }

Auth: admin + finance. Uses the logged-in user's email as the share target.
Upload:
  1. python-docx builds the .docx in-memory using Recast's standard brief layout
  2. Drive API v3 uploads + converts to Google Doc
  3. Drive API shares with the caller's email (writer)
  4. Returns {ok, url}
"""
from http.server import BaseHTTPRequestHandler
from datetime import datetime
import io
import json
import os
import tempfile

from api._shared import (
    json_response, read_body, require_auth, cors_headers,
    SUPABASE_URL, SUPABASE_SERVICE_KEY,
)


# ── Helpers ────────────────────────────────────────────────────────────────

def _fetch_creators(ids):
    """Read creator rows from Supabase using service role key."""
    if not SUPABASE_URL or not SUPABASE_SERVICE_KEY or not ids:
        return []
    import requests as _rq
    # PostgREST 'in' filter syntax: id=in.(uuid1,uuid2,...)
    quoted = ",".join(f'"{i}"' for i in ids)
    try:
        r = _rq.get(
            f"{SUPABASE_URL}/rest/v1/creators",
            params={
                "id": f"in.({quoted})",
                "select": (
                    "id,name,twitch_handle,kick_handle,country,tier,category,"
                    "twitch_30d_ccv,kick_30d_ccv,socials,outreach_status,signed"
                ),
            },
            headers={
                "apikey": SUPABASE_SERVICE_KEY,
                "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
            },
            timeout=10,
        )
        if r.status_code == 200:
            return r.json() or []
        print(f"[export_brief] supabase fetch {r.status_code}: {r.text[:200]}")
    except Exception as e:
        print(f"[export_brief] supabase fetch failed: {e}")
    return []


def _guess_platform(c):
    if c.get("kick_handle") and c.get("twitch_handle"):
        return "Twitch + Kick"
    if c.get("kick_handle"):
        return "Kick"
    if c.get("twitch_handle"):
        return "Twitch"
    return "—"


def _best_ccv_value(c):
    """Return best CCV as an int (for tier bucketing), or 0 if unknown."""
    t = c.get("twitch_30d_ccv") or 0
    k = c.get("kick_30d_ccv") or 0
    return int(max(t, k))


def _best_ccv_label(c):
    v = _best_ccv_value(c)
    return f"{v:,}" if v else "—"


def _followers_label(c):
    """Followers aren't tracked yet; placeholder for manual edit in the Doc."""
    return "~TBC"


def _tier_for(ccv):
    if ccv >= 3000:
        return "TOP"
    if ccv >= 1500:
        return "MID"
    return "BASE"


TIER_LABEL = {
    "TOP": "TOP TIER",
    "MID": "MID TIER",
    "BASE": "BASE TIER",
}

TIER_RANGE = {
    "TOP": "3,000+ avg CCV",
    "MID": "1,500 to 3,000 avg CCV",
    "BASE": "under 1,500 avg CCV",
}


def _handle_line(c):
    if c.get("kick_handle"):
        return f"kick.com/{c['kick_handle']}"
    if c.get("twitch_handle"):
        return f"twitch.tv/{c['twitch_handle']}"
    return ""


def _language_from(c):
    # Rough inference — fall back to English
    country = (c.get("country") or "").lower()
    lang_map = {
        "usa": "English", "uk": "English", "canada": "English", "australia": "English",
        "spain": "Spanish", "mexico": "Spanish", "argentina": "Spanish", "brazil": "Portuguese",
        "germany": "German", "france": "French", "russia": "Russian", "poland": "Polish",
        "japan": "Japanese", "korea": "Korean", "china": "Chinese",
    }
    return lang_map.get(country, "English")


def _market_from(c):
    country = c.get("country") or ""
    if country.lower() in ("usa", "uk", "canada", "australia", "new zealand", "ireland"):
        return "English-speaking"
    if country in ("Spain", "Mexico", "Argentina", "Colombia", "Peru"):
        return "LATAM"
    if country:
        return country
    return "Global"


# ── .docx generation ───────────────────────────────────────────────────────

NAVY = RGBColor(0x1F, 0x2A, 0x44)
GREY_DARK = RGBColor(0x50, 0x50, 0x50)
GREY_MID = RGBColor(0x70, 0x70, 0x70)
GREY_LIGHT = RGBColor(0xA0, 0xA0, 0xA0)


def _shade_cell(cell, hex_rgb):
    """Set a solid fill on a table cell via raw XML."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_rgb)
    tc_pr.append(shd)


def _set_cell_text(cell, lines, sizes=None, bolds=None, colors=None, aligns=None):
    """Replace a cell's content with one paragraph per line, with per-line styling."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    # Clear existing paragraphs except the first
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    first = cell.paragraphs[0]
    first.clear()
    for i, text in enumerate(lines):
        p = first if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if aligns and i < len(aligns) and aligns[i]:
            p.alignment = aligns[i]
        run = p.add_run(text)
        if sizes and i < len(sizes) and sizes[i]:
            run.font.size = Pt(sizes[i])
        if bolds and i < len(bolds):
            run.bold = bolds[i]
        if colors and i < len(colors) and colors[i]:
            run.font.color.rgb = colors[i]


def _label_value_cell(cell, label, value, *, label_size=8, value_size=10):
    """A cell with a small grey label over a normal-weight value — the Casino card style."""
    from docx.shared import Pt
    _set_cell_text(
        cell,
        [label, value],
        sizes=[label_size, value_size],
        bolds=[True, False],
        colors=[GREY_MID, None],
    )


def _merge_row(table, row_idx):
    """Merge every cell in `row_idx` horizontally. Returns the merged cell."""
    row_cells = table.rows[row_idx].cells
    merged = row_cells[0]
    for cell in row_cells[1:]:
        merged = merged.merge(cell)
    return merged


def _build_docx(partner, campaign, criteria, month_year, creators):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Narrower page margins so tables have room to breathe
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # Section label above the cover card
    sec = doc.add_paragraph()
    sr = sec.add_run("CREATOR SHORTLIST OVERVIEW")
    sr.bold = True
    sr.font.size = Pt(9)
    sr.font.color.rgb = GREY_MID

    # ── Cover card (2-col table) ───────────────────────────────────
    cover = doc.add_table(rows=1, cols=2)
    cover.autofit = False
    cover_cells = cover.rows[0].cells
    _set_cell_text(
        cover_cells[0],
        [
            f"RECAST  x  {partner.upper()}",
            f"{campaign}  ·  {month_year}",
            "CONFIDENTIAL",
        ],
        sizes=[22, 14, 9],
        bolds=[True, False, False],
        colors=[NAVY, None, GREY_MID],
    )
    _set_cell_text(
        cover_cells[1],
        ["CONFIDENTIAL", "Prepared by Recast  ·  Not for distribution"],
        sizes=[9, 9],
        bolds=[True, False],
        colors=[NAVY, GREY_MID],
        aligns=[WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT],
    )
    _shade_cell(cover_cells[0], "F3F4F6")
    _shade_cell(cover_cells[1], "F3F4F6")

    doc.add_paragraph()

    # ── Criteria box (1-col table) ─────────────────────────────────
    h = doc.add_paragraph()
    hr = h.add_run("BRIEF CRITERIA")
    hr.bold = True
    hr.font.size = Pt(10)
    hr.font.color.rgb = NAVY

    crit_table = doc.add_table(rows=1, cols=1)
    crit_table.autofit = True
    _set_cell_text(
        crit_table.rows[0].cells[0],
        [criteria or "—"],
        sizes=[10],
        bolds=[False],
    )
    _shade_cell(crit_table.rows[0].cells[0], "FAFAFA")

    doc.add_paragraph()

    # ── Tier legend ────────────────────────────────────────────────
    legend = doc.add_table(rows=1, cols=3)
    legend.autofit = False
    for i, tier_key in enumerate(("TOP", "MID", "BASE")):
        cell = legend.rows[0].cells[i]
        _set_cell_text(
            cell,
            [tier_key, TIER_RANGE[tier_key]],
            sizes=[10, 9],
            bolds=[True, False],
            colors=[NAVY, GREY_MID],
        )
        _shade_cell(cell, "F3F4F6")

    doc.add_paragraph()

    # ── Summary table ──────────────────────────────────────────────
    h2 = doc.add_paragraph()
    h2r = h2.add_run("SUMMARY")
    h2r.bold = True
    h2r.font.size = Pt(10)
    h2r.font.color.rgb = NAVY

    st = doc.add_paragraph()
    stp = st.add_run(
        f"{len(creators)} creator{'' if len(creators)==1 else 's'}  ·  {month_year}"
    )
    stp.italic = True
    stp.font.size = Pt(9)
    stp.font.color.rgb = GREY_MID

    # Sort by CCV descending so the summary reads top→base
    ordered = sorted(creators, key=lambda c: _best_ccv_value(c), reverse=True)

    headers = ["CREATOR", "REGION", "MARKET", "PLATFORM", "AVG CCV", "LANGUAGE", "TIER"]
    summary = doc.add_table(rows=1 + len(ordered), cols=len(headers))
    summary.style = "Light Grid Accent 1"

    for i, col in enumerate(headers):
        _set_cell_text(
            summary.rows[0].cells[i],
            [col],
            sizes=[9],
            bolds=[True],
            colors=[NAVY],
        )

    for idx, c in enumerate(ordered, start=1):
        ccv = _best_ccv_value(c)
        tier = _tier_for(ccv)
        values = [
            c.get("name", ""),
            (c.get("country") or "—")[:20],
            _market_from(c),
            _guess_platform(c),
            f"{ccv:,}" if ccv else "—",
            _language_from(c),
            tier,
        ]
        for i, v in enumerate(values):
            _set_cell_text(
                summary.rows[idx].cells[i],
                [v],
                sizes=[9],
                bolds=[False],
            )

    doc.add_paragraph()

    # ── Creator cards, grouped by tier ─────────────────────────────
    by_tier = {"TOP": [], "MID": [], "BASE": []}
    for c in ordered:
        by_tier[_tier_for(_best_ccv_value(c))].append(c)

    for tier_key in ("TOP", "MID", "BASE"):
        group = by_tier[tier_key]
        if not group:
            continue

        # Tier header band (2-col table)
        band = doc.add_table(rows=1, cols=2)
        band.autofit = False
        bcells = band.rows[0].cells
        _set_cell_text(
            bcells[0],
            [f"TIER: {TIER_LABEL[tier_key]}"],
            sizes=[11],
            bolds=[True],
            colors=[NAVY],
        )
        _set_cell_text(
            bcells[1],
            [f"{TIER_RANGE[tier_key]}  ·  {len(group)} creator{'' if len(group)==1 else 's'}"],
            sizes=[9],
            bolds=[False],
            colors=[GREY_MID],
            aligns=[WD_ALIGN_PARAGRAPH.RIGHT],
        )
        _shade_cell(bcells[0], "EEF0F4")
        _shade_cell(bcells[1], "EEF0F4")

        doc.add_paragraph()

        for c in group:
            ccv = _best_ccv_value(c)
            tier = _tier_for(ccv)

            # Creator card: 5 rows × 3 cols
            card = doc.add_table(rows=5, cols=3)
            card.style = "Light Grid Accent 1"
            card.autofit = True

            # Row 0 — Name | LOCATION | PLATFORM
            r0 = card.rows[0].cells
            _set_cell_text(
                r0[0],
                [c.get("name", "")],
                sizes=[14],
                bolds=[True],
                colors=[NAVY],
            )
            _label_value_cell(r0[1], "LOCATION", c.get("country") or "—")
            _label_value_cell(r0[2], "PLATFORM", _guess_platform(c))

            # Row 1 — FOLLOWERS | AVG CCV (30d) | TIER
            r1 = card.rows[1].cells
            _label_value_cell(r1[0], "FOLLOWERS", _followers_label(c))
            _label_value_cell(r1[1], "AVG CCV (30d)", f"{ccv:,}" if ccv else "—")
            _label_value_cell(r1[2], "TIER", TIER_LABEL[tier])

            # Row 2 — merged — PRIMARY CONTENT / LANGUAGE / MARKET one-liner
            r2 = _merge_row(card, 2)
            content = c.get("category") or "—"
            _set_cell_text(
                r2,
                [
                    f"PRIMARY CONTENT  {content}     LANGUAGE  {_language_from(c)}     MARKET  {_market_from(c)}"
                ],
                sizes=[9],
                bolds=[False],
                colors=[GREY_DARK],
            )

            # Row 3 — merged — WHY THIS FITS + placeholder narrative
            r3 = _merge_row(card, 3)
            _set_cell_text(
                r3,
                [
                    "WHY THIS FITS",
                    "[Add narrative here — observations on content fit, audience, "
                    "posting cadence, alignment with brief criteria, and recommendation.]",
                ],
                sizes=[9, 10],
                bolds=[True, False],
                colors=[NAVY, GREY_LIGHT],
            )
            # make the placeholder italic
            try:
                r3.paragraphs[1].runs[0].italic = True
            except (IndexError, AttributeError):
                pass

            # Row 4 — merged — HANDLE
            r4 = _merge_row(card, 4)
            _set_cell_text(
                r4,
                [f"HANDLE  {_handle_line(c) or '—'}"],
                sizes=[9],
                bolds=[False],
                colors=[GREY_DARK],
            )

            doc.add_paragraph()  # spacer between cards

    # ── Data sources footer ────────────────────────────────────────
    footer = doc.add_table(rows=1, cols=1)
    footer.autofit = True
    _set_cell_text(
        footer.rows[0].cells[0],
        [
            "Data sources: StreamsCharts, TwitchTracker, Kick channel pages, Recast internal database. "
            f"CCV figures reflect 30-day average at time of research ({month_year}). "
            "Narrative sections added manually by the Recast team."
        ],
        sizes=[8],
        bolds=[False],
        colors=[GREY_MID],
    )
    _shade_cell(footer.rows[0].cells[0], "FAFAFA")
    # italicise the footer
    try:
        footer.rows[0].cells[0].paragraphs[0].runs[0].italic = True
    except (IndexError, AttributeError):
        pass

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Drive upload + share ───────────────────────────────────────────────────

def _upload_to_drive_and_share(docx_bytes, filename, share_email):
    """Uploads .docx to Drive, converts to Google Doc, shares with email.
    Returns (webViewLink, file_id) or (None, None) on failure."""
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaIoBaseUpload
    from google.oauth2 import service_account

    creds_json = os.environ.get("GOOGLE_CREDENTIALS_JSON", "")
    if not creds_json:
        return None, None

    tmp_path = None
    try:
        tmp = tempfile.NamedTemporaryFile(mode="w", suffix=".json", delete=False)
        tmp.write(creds_json)
        tmp.close()
        tmp_path = tmp.name

        creds = service_account.Credentials.from_service_account_file(
            tmp_path,
            scopes=["https://www.googleapis.com/auth/drive"],
        )
        drive = build("drive", "v3", credentials=creds, cache_discovery=False)

        media = MediaIoBaseUpload(
            io.BytesIO(docx_bytes),
            mimetype=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
            resumable=False,
        )
        file_body = {
            "name": filename,
            # Google Docs MIME type — converts .docx into a native Doc
            "mimeType": "application/vnd.google-apps.document",
        }
        folder_id = os.environ.get("GOOGLE_DRIVE_FOLDER_ID", "").strip()
        if folder_id:
            file_body["parents"] = [folder_id]

        file = (
            drive.files()
            .create(
                body=file_body,
                media_body=media,
                fields="id,webViewLink",
                supportsAllDrives=True,
            )
            .execute()
        )

        file_id = file.get("id")
        link = file.get("webViewLink")

        # Share with the caller's email as writer
        if share_email:
            try:
                drive.permissions().create(
                    fileId=file_id,
                    body={"type": "user", "role": "writer", "emailAddress": share_email},
                    sendNotificationEmail=True,
                    supportsAllDrives=True,
                ).execute()
            except Exception as e:
                print(f"[export_brief] share failed for {share_email}: {e}")

        return link, file_id
    except Exception as e:
        print(f"[export_brief] drive upload failed: {e}")
        return None, None
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
            except Exception:
                pass


# ── Handler ────────────────────────────────────────────────────────────────

class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        user = require_auth(self, required_roles=("admin", "finance"))
        if not user:
            return
        try:
            body = read_body(self)
            partner = (body.get("partner") or "").strip()
            campaign = (body.get("campaign") or "").strip()
            criteria = (body.get("criteria") or "").strip()
            month_year = (body.get("month_year") or "").strip() or datetime.now().strftime("%B %Y")
            creator_ids = body.get("creator_ids") or []

            if not partner or not campaign or not creator_ids:
                json_response(self, 400, {"error": "partner, campaign, creator_ids are required"})
                return

            creators = _fetch_creators(creator_ids)
            if not creators:
                json_response(self, 404, {"error": "no creators found for the provided ids"})
                return

            # Preserve selection order
            by_id = {c["id"]: c for c in creators}
            ordered = [by_id[i] for i in creator_ids if i in by_id]

            docx_bytes = _build_docx(partner, campaign, criteria, month_year, ordered)

            safe_campaign = "".join(ch for ch in campaign if ch.isalnum() or ch in (" ", "-", "_")).strip()
            filename = f"Recast_{partner}_{safe_campaign}_{month_year.replace(' ', '')}"

            share_email = user.get("email") or ""
            url, file_id = _upload_to_drive_and_share(docx_bytes, filename, share_email)

            if not url:
                json_response(self, 500, {"error": "Drive upload failed"})
                return

            json_response(
                self,
                200,
                {
                    "ok": True,
                    "url": url,
                    "file_id": file_id,
                    "shared_with": share_email,
                    "creators": len(ordered),
                },
            )
        except Exception as e:
            json_response(self, 500, {"error": str(e)})

    def do_OPTIONS(self):
        self.send_response(204)
        cors_headers(self)
        self.end_headers()
